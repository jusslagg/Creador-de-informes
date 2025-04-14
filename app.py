import streamlit as st
import pandas as pd
import google.generativeai as genai
from docx import Document
from io import BytesIO
import re
import requests
from bs4 import BeautifulSoup
import PyPDF2
import os

# Configuración inicial de la API de Gemini
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-pro-latest')

# Validar URLs
def is_valid_url(url):
    regex = re.compile(
        r'^(?:http|ftp)s?://' 
        r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+(?:[A-Z]{2,6}\.?|[A-Z0-9-]{2,}\.?)|'
        r'localhost|'  
        r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})' 
        r'(?::\d+)?'  
        r'(?:/?|[/?]\S+)$', re.IGNORECASE)
    return re.match(regex, url) is not None

# Lectura de archivos
def read_excel(file):
    try:
        return pd.read_excel(file, engine='openpyxl'), "excel/csv"
    except Exception as e:
        st.error(f"Error al leer Excel: {e}")
        return pd.DataFrame(), "none"

def read_csv(file):
    try:
        return pd.read_csv(file, encoding='latin1', on_bad_lines='skip'), "excel/csv"
    except Exception as e:
        st.error(f"Error al leer CSV: {e}")
        return pd.DataFrame(), "none"

def read_docx(file):
    document = Document(file)
    text = "\n".join([paragraph.text for paragraph in document.paragraphs])
    return pd.DataFrame([text], columns=['text']), "word"

def read_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        return pd.DataFrame([text], columns=['text']), "pdf"
    except Exception as e:
        st.error(f"Error al leer PDF: {e}")
        return pd.DataFrame(), "none"

def fetch_web_content(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        text = soup.get_text(separator='\n')
        return pd.DataFrame([text], columns=['text']), "web"
    except Exception as e:
        st.error(f"Error al leer URL: {e}")
        return pd.DataFrame(), "none"

# Calcular cuartiles
def calculate_quartiles(df, column):
    if column in df.columns and pd.api.types.is_numeric_dtype(df[column]):
        quartiles = pd.qcut(df[column], q=4, labels=["Q4 (Peor)", "Q3", "Q2", "Q1 (Mejor)"], duplicates='drop')
        return df.assign(Quartil=quartiles)
    else:
        st.error(f"La columna '{column}' no es válida o no es numérica.")
        return df

# PROMPT optimizado
def generate_prompt(level, context_text, df_display):
    data_string = df_display.to_string(index=False)
    prompts = {
        "Análisis de Ranking": f"""
Actúa como un analista senior en control de gestión de un call center. Tu tarea es realizar un análisis exhaustivo del rendimiento de los agentes basado en los datos que te proporciono.

Contexto general del análisis:
{context_text}

Datos a analizar:
{data_string}

Realiza lo siguiente:
1. Detecta fortalezas y debilidades por agente.
2. Compara el rendimiento entre los mejores y peores.
3. Identifica tendencias generales o anomalías en los datos.
4. Extrae insights clave a nivel operativo.
5. Genera un ranking Top 10 (mejores y peores).
6. Analiza los cuartiles calculados: Q1 (Mejor 25%) a Q4 (Peor 25%).
7. Propón recomendaciones concretas para mejorar el desempeño general.
""",
        "Tiempos Productivos, Hold, Baño, Break": f"""
Eres un analista experto en eficiencia operativa. Evalúa el uso del tiempo por parte de los agentes de un call center.

Contexto del análisis:
{context_text}

Datos proporcionados:
{data_string}

Tu análisis debe incluir:
1. Evaluación de tiempo productivo vs improductivo.
2. Análisis de duración y frecuencia de pausas (hold, baño, break).
3. Detección de patrones de improductividad o abuso de pausas.
4. Comparación entre agentes.
5. Recomendaciones claras para mejorar la gestión del tiempo y productividad.
""",
        "Tableros de Incidencias": f"""
Actúa como especialista en gestión de calidad. Tu objetivo es analizar las incidencias reportadas en el call center.

Contexto:
{context_text}

Datos disponibles:
{data_string}

Tareas a realizar:
1. Identifica las incidencias más frecuentes y su impacto.
2. Detecta agentes o equipos con más incidencias.
3. Evalúa la efectividad de los procesos de resolución.
4. Propón mejoras de procesos y prevención de incidentes futuros.
""",
        "Satisfacción del Cliente": f"""
Eres experto en experiencia del cliente. Analiza la satisfacción según métricas como NPS, CSAT, y comentarios.

Contexto del análisis:
{context_text}

Datos recibidos:
{data_string}

Tu informe debe incluir:
1. Análisis de tendencias en satisfacción.
2. Identificación de factores clave que afectan la percepción del cliente.
3. Segmentación si hay diferentes perfiles de clientes.
4. Propuestas de mejora específicas basadas en los hallazgos.
""",
        "Costos y Rentabilidad": f"""
Actúa como analista financiero de un call center. Evalúa los costos operativos y márgenes de rentabilidad.

Contexto del análisis:
{context_text}

Datos:
{data_string}

Puntos clave a desarrollar:
1. Desglose de los costos operativos y su evolución.
2. Análisis de eficiencia: costos vs productividad.
3. Identificación de oportunidades de ahorro.
4. Recomendaciones para mejorar la rentabilidad sin afectar el servicio.
""",
        "Libre": f"""
Actúa como consultor experto. Tienes total libertad para extraer y explicar los hallazgos más importantes a partir del contexto y los datos.

Contexto general:
{context_text}

Datos:
{data_string}

Tu informe debe incluir:
1. Análisis profundo y claro de lo que encuentres relevante.
2. Insights accionables.
3. Recomendaciones estratégicas bien fundamentadas.
"""
    }
    return prompts[level]

# Interfaz de usuario
uploaded_file = st.file_uploader("Carga tu archivo Excel, CSV, Word o PDF", type=["xls", "xlsx", "csv", "docx", "pdf"])
web_url = st.text_input("Ingresa la URL de la Página Web")
context_text = st.text_area("Ingresa el contexto para el análisis")

df, data_type = pd.DataFrame(), "none"
if uploaded_file:
    file_type = uploaded_file.type
    if file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        df, data_type = read_excel(uploaded_file)
    elif file_type == "text/csv":
        df, data_type = read_csv(uploaded_file)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        df, data_type = read_docx(uploaded_file)
    elif file_type == "application/pdf":
        df, data_type = read_pdf(uploaded_file)
    else:
        st.error("Tipo de archivo no soportado.")
elif web_url and is_valid_url(web_url):
    df, data_type = fetch_web_content(web_url)

# Mostrar datos
if data_type != "none" and not df.empty:
    st.write("Archivo cargado exitosamente!")
    for col in df.columns:
        if df[col].dtype == 'object':
            df[col] = df[col].astype(str)
    st.dataframe(df)

    selected_columns = st.multiselect("Selecciona las columnas a analizar", df.columns.tolist())
    if selected_columns:
        df_with_quartiles = calculate_quartiles(df, selected_columns[0])
        st.dataframe(df_with_quartiles)
        df_display = df_with_quartiles.head(100)
    else:
        df_display = df.head(100)

    level = st.selectbox("Selecciona el nivel del análisis", [
        "Análisis de Ranking", 
        "Tiempos Productivos, Hold, Baño, Break", 
        "Tableros de Incidencias", 
        "Satisfacción del Cliente", 
        "Costos y Rentabilidad", 
        "Libre"
    ])

    if st.button("Empezar el análisis 🚀"):
        if df.empty:
            st.error("No hay datos para analizar.")
        else:
            prompt = generate_prompt(level, context_text, df_display)
            try:
                contents = [
                    {"role": "user", "parts": ["Eres un analista senior. Aplica un análisis riguroso y multifacético."]},
                    {"role": "user", "parts": [prompt]}
                ]
                response = model.generate_content(contents)
                informe = response.text
                st.write("--- Informe Generado ---")
                for line in informe.splitlines():
                    if line.startswith(("Título principal:", "Subtítulo")):
                        st.markdown(f"<h2 style='color: blue;'>{line}</h2>", unsafe_allow_html=True)
                    else:
                        st.write(line)

                document = Document()
                document.add_heading('Informe Generado por CAT-AI', 0)
                for line in informe.splitlines():
                    document.add_paragraph(line)
                docx_stream = BytesIO()
                document.save(docx_stream)
                docx_stream.seek(0)
                st.download_button(
                    label="Descargar informe en Word",
                    data=docx_stream,
                    file_name="informe.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Error al generar el informe: {e}")
else:
    st.info("Por favor, carga un archivo o ingresa una URL válida.")
